from flask import flash, render_template, redirect, url_for, Blueprint, g, request, session, send_file, send_from_directory
from flask import current_app as app
from flask_wtf import FlaskForm
from flask_wtf.recaptcha import RecaptchaField
from wtforms import EmailField, StringField, PasswordField, validators

from sqlalchemy import select, delete

import requests, datetime

from email_validator import validate_email, EmailNotValidError
import phonenumbers

from .db_IXIPC import get_session
from .models import User, Abstract, AbstractType, Author, AbstractAuthor, Institution, Affiliation, Payment, PaymentMethod, PaymentStatus
from ..files import UploadedFile

import json, functools
from hashlib import sha256
import hmac
from secrets import token_hex

import csv
from io import BytesIO, StringIO
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from docx import Document
from docx.shared import Pt

bp = Blueprint('IX_IPC', __name__, template_folder='templates')


abstract_submission_open = True


def login_IXIPC_required(view):
    """Guarantees the user is logged in
     
    Decorator that guarantees the user is logged in, redirecting to the main 
    page if they are not.
    """

    @functools.wraps(view)
    def wrapped_view(**kwargs):
        if g.IXIPC_user is None:
            if 'language' not in kwargs:
                kwargs['language'] = 'pt'
            
            return redirect(
                url_for('IX_IPC.IXIPC', language=kwargs['language'])
            )

        return view(**kwargs)

    return wrapped_view

def login_IXIPC_management_required(view):
    """Guarantees a manager is logged in
     
    Decorator that guarantees a manager is logged in, redirecting to the main 
    page if they are not.
    """

    @functools.wraps(view)
    def wrapped_view(**kwargs):
        if g.IXIPC_user is None or (not g.IXIPC_user.organizer):
            if 'language' not in kwargs:
                kwargs['language'] = 'pt'
            
            return redirect(
                url_for('IX_IPC.IXIPC', language=kwargs['language'])
            )

        return view(**kwargs)

    return wrapped_view


class RegistrationForm(FlaskForm):
    name = StringField('registration-form-name', [validators.DataRequired(), validators.Length(max=50)])
    email = StringField('registration-form-email', [validators.DataRequired(), validators.Email()])
    recaptcha = RecaptchaField()

class LoginForm(FlaskForm):
    email = StringField('login-form-email', [validators.DataRequired(), validators.Email()])
    password = PasswordField('login-form-password', [validators.DataRequired()])
    recaptcha = RecaptchaField()

class ManagementLoginForm(FlaskForm):
    password = PasswordField('login-form-password', [validators.DataRequired()])


class BankDetails():
    def __init__(self) -> None:
        self.beneficiario = app.config['IXIPC_BANK_BENEFICIARY']
        self.iban = app.config['IXIPC_BANK_IBAN']
        self.bic = app.config['IXIPC_BANK_BIC']

@bp.route('/IX_IPC/<language:language>')
@bp.route('/IX_IPC')
@bp.route('/IX_Iberian_Primatological_Conference/<language:language>')
@bp.route('/IX_Iberian_Primatological_Conference')
def IXIPC(language='pt'):
    """IXIPC main page"""

    g.links[2]['active'] = True

    user_id = None
    with get_session() as db_session:
        abstracts = []
        payment_status = 0
        organizer = False
        if g.IXIPC_user:
            abstracts = db_session.execute(
                select(Abstract).filter_by(owner=g.IXIPC_user.id)
            ).scalars()

            user_id = g.IXIPC_user.id

            if g.IXIPC_user.paid_registration:
                payment_status = 1
            elif g.IXIPC_user.payment_id:
                payment_status = 2

            if g.IXIPC_user.organizer:
                organizer = True
        
        return render_template(
            'IX_IPC.html',
            lang=language,
            registration_form=RegistrationForm(),
            login_form=LoginForm(),
            text_column=True,
            light_background=True,
            abstracts=abstracts,
            site_map=True,
            payment_status=payment_status,
            organizer=organizer,
            submission_open=abstract_submission_open,
            bank_details=BankDetails(),
            user_id=user_id
        )


def sanitize_email(email: str) -> str|None:
    """Helper function to sanitize input emails"""

    try:
        emailinfo = validate_email(email, check_deliverability=False)
        email = emailinfo.normalized
    except EmailNotValidError as e:
        return None
    return email


@bp.route('/IX_IPC/register/<language:language>', methods=['POST'])
def register_user(language):
    """Registers a new user
    
    This route registers a new user and logs them in, redirecting to the main
    page with a welcoming message. In the registration process, an email is sent 
    to the user informing their credentials. If anything goes wrong (invalid 
    email, the email is already registered, etc) a warning is displayed.
    """

    g.links[2]['active'] = True

    form = RegistrationForm()
    valid = form.validate_on_submit()
    if not valid:
        flash('invalid captcha!', 'warning')
        return redirect(url_for("IX_IPC.IXIPC", language=language) + "#alerts")
    
    name = form.name.data.strip()
    email = form.email.data

    if not email:
        flash('Email is required.', 'warning')
    elif not name:
        flash('Name is required.', 'warning')
    elif valid:
        email = sanitize_email(email)
        if email:
            with get_session() as db_session:
                try:
                    password = token_hex(16)
                    user = User(name, email, password=password)
                    db_session.add(user)
                    db_session.commit()
                    session['IXIPC_user_id'] = user.id
                except:
                    url = url_for(
                        'IX_IPC.recover_credentials', 
                        language=language, 
                        email=email
                    )
                    flash(
                        app.translate(
                          'IXIPC-email-already-registered', 
                          language, {
                              'email': email, 
                              'recover_credentials': url
                          }
                        ), 'warning'
                    )
                else:
                    try:
                        app.send_email(
                            app.translate('IXIPC-welcome-email-subject', language),
                            app.translate(
                                'IXIPC-welcome-email-body', 
                                language, {
                                'email': email,
                                'name': name,
                                'password': password
                            }),
                            [email]
                        )
                    except:
                        flash(
                            app.translate(
                              'IXIPC-email-failed', language, {'email': email}
                            ), 'warning'
                        )
                    else:
                        flash(
                            app.translate(
                                'IXIPC-account-creation-successful', language
                            ), 'success'
                        )
                        return redirect(url_for("IX_IPC.IXIPC", language=language) + "#alerts")
        else:
            flash(app.translate('IXIPC-invalid-email', language), 'warning')

    return redirect(url_for("IX_IPC.IXIPC", language=language) + "#alerts")


@bp.route('/IX_IPC/recover_credentials/<language:language>/<string:email>')
def recover_credentials(language, email):
    """Resends the email with the credentials"""

    email = sanitize_email(email)
    if email:
        with get_session() as db_session:
            user = db_session.execute(select(User).filter_by(email=email)).scalar_one()
            password = token_hex(16)
            user.update_password(password)
            db_session.commit()

            app.send_email(
                app.translate('IXIPC-recover-credentials-email-subject', language),
                app.translate(
                    'IXIPC-recover-credentials-email-body', 
                    language, {
                    'email': email,
                    'name': user.name,
                    'password': password
                }),
                [email]
            )
            flash(
                app.translate(
                    'IXIPC-recover-credentials-notification', 
                    language
                ), 'success'
            )
    else:
        flash(app.translate('IXIPC-invalid-email', language), 'warning')
    return redirect(url_for("IX_IPC.IXIPC", language=language) + '#alerts')


@bp.route('/IX_IPC/login/<language:language>', methods=['POST'])
def login(language):
    """Logs a user in
    
    Logs a user in and redirects to the main page. Displays a warning in case
    the login is unsuccessful. 
    """

    form = LoginForm()
    if form.validate_on_submit():
        email = sanitize_email(form.email.data)

        try:
            with get_session() as db_session:
                user = db_session.execute(
                    select(User).filter_by(email=email)
                ).scalar_one()

            if user and user.check_password(form.password.data):
                session.clear()
                session['IXIPC_user_id'] = user.id
                if user.organizer:
                    session['IXIPC_manager'] = True
            else:
                flash(
                    app.translate('IXIPC-login-wrong-email-or-password', language), 
                    'warning'
                )
                return redirect(url_for('IX_IPC.IXIPC', language=language) + '#alerts')
        except:
            flash(
                app.translate('IXIPC-login-wrong-email-or-password', language), 
                'warning'
            )
            return redirect(url_for('IX_IPC.IXIPC', language=language) + '#alerts')
    else:
        flash('Invalid CAPTCHA', 'warning')
        return redirect(url_for('IX_IPC.IXIPC', language=language) + '#alerts')
    
    return redirect(url_for('IX_IPC.IXIPC', language=language))


@bp.route('/IX_IPC/logout/<language:language>', methods=['POST'])
@login_IXIPC_required
def logout(language):
    """Logs a user out
    
    Logs a user out and redirects to the main page.
    """

    session.clear()
    return redirect(url_for('IX_IPC.IXIPC', language=language))


@bp.before_app_request
def load_logged_in_IXIPC_user() -> None:
    """Loads the current user
    
    Function that runs before every request and adds the current user to g.
    """

    user_IXIPC_id = session.get('IXIPC_user_id')
    if user_IXIPC_id is None:
        g.IXIPC_user = None
    else:
        try:
            with get_session() as db_session:
                g.IXIPC_user = db_session.get(User, user_IXIPC_id)
        except:
            g.IXIPC_user = None

    # Check if the user is a manager
    if not session.get('IXIPC_manager'):
        g.IXIPC_manager = False
    else:
        g.IXIPC_manager = True


@bp.route('/IX_IPC/save_personal_data', methods=['POST'])
@login_IXIPC_required
def save_personal_data():
    """Saves current user's data to the database"""

    with get_session() as db_session:
        user = db_session.get(User, g.IXIPC_user.id)
        user.first_name = request.form['first-name'].strip()
        user.last_name = request.form['last-name'].strip()
        user.institution = request.form['institution'].strip()
        user.student = request.form['student'] == 'true'
        user.member = request.form['member'] == 'true'
        user.scholarship = request.form['scholarship'] == 'true'
        user.unemployed = request.form['unemployed'] == 'true'
        user.competition_talk = request.form['competition_talk'] == 'true'
        user.competition_photography = request.form['competition_photography'] == 'true'
        # user.dinner = request.form['dinner'] == 'true'
        # user.dinner_type = request.form['dinner_type']

        db_session.commit()

        return json.dumps({}), 200


@bp.route('/IX_IPC/create_abstract/<language:language>', methods=['POST'])
@login_IXIPC_required
def create_new_abstract(language='pt'):
    """Creates and returns a new abstract"""

    id = int(json.loads(save_abstract_local(None, g))['id'])
    return load_closed_abstract(id, language), 200


@bp.route('/IX_IPC/closed_abstract/<language:language>', methods=['POST'])
@login_IXIPC_required
def closed_abstract(language='pt'):
    """Returns the closed form of an existing abstract"""

    id = int(request.form['abstract-id'])
    with get_session() as db_session:
        abstract = db_session.get(Abstract, id)

        if abstract.owner == g.IXIPC_user.id:
            return load_closed_abstract(id, language), 200
        else:
            return json.dumps({'error': 'access unauthorized'}), 401


def load_closed_abstract(abstract: int|Abstract, language) -> str:
    """Returns a closed abstract HTML element
    
    Receives an Abstract object or an index to an abstract and return a json 
    string containing the abstract id and the HTML element for the closed 
    abstract.
    """

    if isinstance(abstract, int):
        with get_session() as db_session:
            abstract = db_session.get(Abstract, abstract)
    
    return json.dumps({
        'id': abstract.id,
        'html': render_template(
            'abstract-form-closed.html',
            abstract = abstract,
            csrf_token = request.form['csrf_token'],
            reload=True,
            lang=language
        )})


@bp.route('/IX_IPC/delete_abstract', methods=['POST'])
@login_IXIPC_required
def delete_abstract():
    """Deletes an existing abstract
    
    Deletes the abstract data as well as the corresponding abstract authors.
    """

    with get_session() as db_session:
        id = request.form['abstract-id']

        abstract = db_session.get(Abstract, id)
        delete_abstract_authors = delete(AbstractAuthor)\
                                    .where(AbstractAuthor.abstract_id == id)
        
        if g.IXIPC_user.id == abstract.owner:
            try:
                db_session.delete(abstract)
                db_session.execute(delete_abstract_authors)
                db_session.commit()
            except:
                return json.dumps({'error': 'error deleting abstract'}), 500
            return json.dumps({}), 204
        return json.dumps({'error': 'access unauthorized'}), 401


def save_abstract_local(form, g) -> str:
    """Saves an abstract to the database
    
    Helper function that saves an abstract to the database, creating a new
    abstract if an abstract id is not provided in the request.
    """

    with get_session() as db_session:
        abstract = None
        if form and 'abstract-id' in form and len(form['abstract-id']) > 0:
            id = int(form['abstract-id'])
            abstract = db_session.get(Abstract, id)

            if abstract.owner != g.IXIPC_user.id:
                return json.dumps({'error': 'access unauthorized'})
        else:
            abstract = Abstract(owner=g.IXIPC_user.id)
            db_session.add(abstract)

        if form:
            if 'title' in form:
                abstract.title = form['title'].strip()
                # TODO: update this rejection criteria
                # if len(abstract.title) > 150:
                #     return json.dumps({'error': 'title too long'})

            if 'abstract-body' in form:
                abstract.abstract = form['abstract-body'].strip()
                # TODO: update this rejection criteria
                # if len(abstract.abstract) > 500:
                #     return json.dumps({'error': 'abstract too long'})

            if 'abstract_type' in form:
                abstract_type = form['abstract_type']
                if abstract_type == 'poster':
                    abstract.abstract_type = AbstractType.POSTER
                elif abstract_type == 'presentation':
                    abstract.abstract_type = AbstractType.PRESENTATION
                else:
                    return json.dumps({'error': 'unrecognized abstract type'})
            
            if 'abstract-keywords' in form:
                # TODO: process keywords to guarantee they are shown in a uniform fashion
                abstract.keywords = form['abstract-keywords'].strip()

            if 'scientific-area' in form:
                if form['scientific-area'] == 'other':
                    abstract.scientific_area = form['scientific-area-desc']
                else:
                    abstract.scientific_area = form['scientific-area']
        else:
            abstract.title = ''
            abstract.abstract = ''
            abstract.abstract_type = AbstractType.POSTER

        db_session.commit()

        return json.dumps({'id': abstract.id})


@bp.route('/IX_IPC/load_abstract/<language:language>/', methods=['POST'])
@login_IXIPC_required
def load_abstract(language, id=None, csrf_token = None):
    """Loads an abstract in the open form
    
    Returns a json string containing the abstract id and an HTML element 
    representing the open abstract in the submitted or unsubmitted form 
    depending on submission status.
    """

    if not id:
        id = request.form['abstract-id']
    
    if not csrf_token:
        csrf_token = request.form['csrf_token']

    with get_session() as db_session:
        abstract = db_session.get(Abstract, id)

        abstract.authors, abstract.affiliations = process_authors_affiliations(abstract.id)
        for author in abstract.authors:
            author['affiliations'] = map(lambda x: str(x+1), author['affiliations'])

        if g.IXIPC_user.id == abstract.owner:
            if abstract.submitted:
                return json.dumps({
                  'id': abstract.id,
                  'html': render_template(
                      'abstract-form-open-submitted.html', 
                      lang=language, 
                      form=request.form, 
                      reload=True,
                      abstract=abstract,
                      csrf_token=csrf_token,
                      authors=get_abstract_authors_list(id)
                  )
              }), 200
            else:
              return json.dumps({
                  'id': abstract.id,
                  'html': render_template(
                      'abstract-form-open-unsubmitted.html', 
                      lang=language, 
                      form=request.form, 
                      reload=True,
                      abstract=abstract,
                      csrf_token=csrf_token,
                      submission_open=abstract_submission_open,
                      conditionally_accepted=(abstract.acceptance_status==3)
                  )
              }), 200
        else:
            return json.dumps({'error': 'access unauthorized'}), 401


@bp.route('/IX_IPC/save_abstract', methods=['POST'])
@login_IXIPC_required
def save_abstract():
    """Saves an abstract"""

    return save_abstract_local(request.form, g), 200


@bp.route('/IXIPC/new_author', methods=['POST'])
@login_IXIPC_required
def new_author():
    """Creates a new author"""

    with get_session() as db_session:
        author = Author(created_by=g.IXIPC_user.id)
        db_session.add(author)
        db_session.commit()

        return json.dumps({'id': author.id}), 200


@bp.route('/IX_IPC/load_authors', methods=['POST'])
@login_IXIPC_required
def load_authors():
    """Loads authors available to current user"""

    with get_session() as db_session:
        authors = db_session.execute(
            select(Author).where(Author.created_by == g.IXIPC_user.id)
        ).scalars()

        authors_list = []
        for author in authors:
            first_name = author.first_name if author.first_name else '' 
            last_name = author.last_name if author.last_name else '' 
            authors_list.append({
                'firstName': first_name,
                'lastName': last_name,
                'id': author.id
            })

        return json.dumps({'authors': authors_list}), 200


@bp.route('/IX_IPC/save_authors', methods=['POST'])
@login_IXIPC_required
def save_authors():
    """Saves a list of authors
    
    Saves a list of authors provided in the request.
    """

    authors = json.loads(request.form['authors'])
    with get_session() as db_session:
        authors_list = []
        for (id, author_new) in authors.items():
            author = db_session.get(Author, int(id))
            if g.IXIPC_user.id == author.created_by:
                author.first_name = author_new['firstName'].strip()
                author.last_name = author_new['lastName'].strip()
                authors_list.append(author)
        
        db_session.add_all(authors_list)
        db_session.commit()
    
    return json.dumps({}), 200


@bp.route('/IX_IPC/save_affiliations', methods=['POST'])
@login_IXIPC_required
def save_affiliations():
    """Saves authors' affiliations
    
    Saves the affiliations for a list of authors provided in the request.
    """

    affiliations = json.loads(request.form['affiliations'])
    print(affiliations)
    with get_session() as db_session:
        new_affiliations = []
        for (author_id, institution) in affiliations.items():
            db_session.execute(delete(Affiliation).where(
                Affiliation.author_id == author_id
            ))
            for (institution_id, order) in institution:
                affiliation = Affiliation(
                    author_id=author_id, 
                    institution_id=institution_id, 
                    order=order
                )
                new_affiliations.append(affiliation)

        db_session.add_all(new_affiliations)
        db_session.commit()

    return json.dumps({}), 200


def get_affiliations(author_id: int) -> list[object]|None:
    """Returns a list of authors' affiliations

    Returns a list of authors' affiliations for a given author if or None if
    the author owner is not the current user.
    """

    with get_session() as db_session:
        if db_session.get(Author, author_id).created_by == g.IXIPC_user.id:
            query = select(Affiliation)\
              .where(Affiliation.author_id == author_id)\
              .order_by(Affiliation.order)
            
            affiliations_list = db_session.execute(query).scalars()

            affiliations = []
            for affiliation in affiliations_list:
                affiliations.append({
                    'authorId': author_id,
                    'institutionId': affiliation.institution_id,
                    'order': affiliation.order,
                })
            
            return affiliations
        else:
            return None


@bp.route('/IX_IPC/load_affiliations', methods=['POST'])
@login_IXIPC_required
def load_affiliations():
    """Loads authors' affiliations
    
    Loads the authors' affiliations in request.
    """
    
    author_id = json.loads(request.form['authorId'])
    return json.dumps({'affiliations': get_affiliations(author_id)}), 200


@bp.route('/IX_IPC/save_abstract_authors', methods=['POST'])
@login_IXIPC_required
def save_abstract_authors():
    """Saves the abstract authors
    
    Saves the abstract authors for a list of abstracts provided in the request.
    """

    abstracts = json.loads(request.form['abstractAuthors'])
    with get_session() as db_session:
        abstract_authors = []
        for (abstract_id, abstract) in abstracts.items():
            if db_session.get(Abstract, abstract_id).owner == g.IXIPC_user.id:
                db_session.execute(delete(AbstractAuthor).where(
                    AbstractAuthor.abstract_id == abstract_id
                ))
                for (author_id, order, presenter) in abstract:
                    abstract_author = AbstractAuthor(
                        author_id=author_id, 
                        abstract_id=abstract_id, 
                        order=order, 
                        presenter=presenter
                    )
                    abstract_authors.append(abstract_author)
                
        db_session.add_all(abstract_authors)
        db_session.commit()
    
    return json.dumps({}), 200


def get_abstract_authors_list(abstract_id: int) -> list[object]|None:
    """Returns a list of abstract authors

    Returns a list of the abstract authors for a given abstract id or None if
    the abstract owner is not the current user.
    """

    with get_session() as db_session:
        if db_session.get(Abstract, abstract_id).owner == g.IXIPC_user.id:
            query = select(
                  AbstractAuthor.abstract_id, AbstractAuthor.author_id, 
                  AbstractAuthor.presenter, AbstractAuthor.order, 
                  Author.first_name, Author.last_name
              ).select_from(AbstractAuthor)\
              .join(Author, AbstractAuthor.author_id == Author.id)\
              .where(AbstractAuthor.abstract_id == abstract_id)\
              .order_by(AbstractAuthor.order)
            
            abstract_authors_list = db_session.execute(query)

            abstract_authors = []
            for abstract_author in abstract_authors_list:
                abstract_authors.append({
                    'authorId': abstract_author.author_id,
                    'firstName': abstract_author.first_name,
                    'lastName': abstract_author.last_name,
                    'order': abstract_author.order,
                    'presenter': abstract_author.presenter
                })
            
            return abstract_authors
        else:
            return None


@bp.route('/IX_IPC/load_abstract_authors', methods=['POST'])
@login_IXIPC_required
def load_abstract_authors():
    """Loads abstract authors
    
    Loads the abstract authors for the abstract in request.
    """
    
    abstract_id = json.loads(request.form['abstractId'])
    return json.dumps({'authors': get_abstract_authors_list(abstract_id)}), 200


@bp.route('/IX_IPC/submit_abstract/<language:language>', methods=['POST'])
@login_IXIPC_required
def submit_abstract(language):
    """Submits an abstract
    
    Submits the abstract and returns the corresponding submitted abstract HTML 
    element or an error message.
    """
    
    id = request.form['abstract-id']
    output = json.loads(save_abstract_local(request.form, g))

    if 'error' not in output:
        with get_session() as db_session:
            abstract = db_session.get(Abstract, id)
            abstract.submit()
            db_session.commit()

            return load_abstract(language=language, id=id)
    else:
        return output, 500


@bp.route('/IXIPC/new_institution', methods=['POST'])
@login_IXIPC_required
def new_institution():
    """Creates a new institution"""

    with get_session() as db_session:
        institution = Institution(created_by=g.IXIPC_user.id)
        db_session.add(institution)
        db_session.commit()

        return json.dumps({'id': institution.id}), 200


@bp.route('/IX_IPC/load_institution', methods=['POST'])
@login_IXIPC_required
def load_institution():
    """Loads an institution"""

    id = json.loads(request.form['id'])
    with get_session() as db_session:
        institution = db_session.get(Institution, id)

        if institution.created_by == g.IXIPC_user.id:
            data = {}
            data['id'] = id
            data['name'] = institution.name if institution.name else ''
            data['address'] = institution.address if institution.address else ''
            data['country'] = institution.country if institution.country else ''
            return json.dumps({'institution': data}), 200
        else:
            return json.dumps({'error': 'access unauthorized'}), 401


@bp.route('/IX_IPC/load_institutions', methods=['POST'])
@login_IXIPC_required
def load_institutions():
    """Loads institutions available to current user"""

    with get_session() as db_session:
        institutions = db_session.execute(
            select(Institution).where(Institution.created_by == g.IXIPC_user.id)
        ).scalars()

        institutions_list = []
        for institution in institutions:
            name = institution.name if institution.name else '' 
            address = institution.address if institution.address else '' 
            country = institution.country if institution.country else '' 
            institutions_list.append({
                'name': name,
                'address': address,
                'country': country,
                'id': institution.id
            })

        return json.dumps({'institutions': institutions_list}), 200

@bp.route('/IX_IPC/load_all_affiliations', methods=['POST'])
@login_IXIPC_required
def load_all_affiliations():
    """Loads all affiliations of the current user's authors"""

    with get_session() as db_session:
        affiliations = db_session.execute(
            select(Affiliation)
            .select_from(Affiliation)
            .join(Author, Affiliation.author_id == Author.id)
            .where(Author.created_by == g.IXIPC_user.id)
            .order_by(Affiliation.order)
        ).scalars()

        affiliations_list = []
        for affiliation in affiliations:
            affiliations_list.append({
                'author_id': affiliation.author_id,
                'institution_id': affiliation.institution_id,
                'order': affiliation.order
            })

        return json.dumps({'affiliations': affiliations_list}), 200


@bp.route('/IX_IPC/save_institutions', methods=['POST'])
@login_IXIPC_required
def save_institutions():
    """Saves a list of institutions
    
    Saves the list of institutions provided in the request.
    """

    institutions = json.loads(request.form['institutions'])
    with get_session() as db_session:
        institutions_list = []
        for (id, institution_new) in institutions.items():
            institution = db_session.get(Institution, int(id))
            if g.IXIPC_user.id == institution.created_by:
                institution.name = institution_new['name'].strip()
                institution.address = institution_new['address'].strip()
                institution.country = institution_new['country'].strip()
                institutions_list.append(institution)
        
        db_session.add_all(institutions_list)
        db_session.commit()
    
    return json.dumps({}), 200


# https://helpdesk.ifthenpay.com/en/support/solutions/articles/79000141345-api-mbway-v2-rest
# If then pay API request
# {
#     "mbWayKey": "MBWAY_KEY", //(string) mandatory (assigned by ifthenpay)
#     "orderId": "ORDER_ID", //(string) mandatory (15 chars max)
#     "amount": "AMOUNT", //(string) mandatory (decimal separator ".")
#     "mobileNumber": "MOBILE_NUMBER", //(string) mandatory
#     "email": "EMAIL", //(string) optional (100 chars max)
#     "description": "DESCRIPTION", //(string) optional (100 chars max)
# }
#
# and response:
# {
#     "Amount": "33.61",
#     "Message": "Pending",
#     "OrderId": "1887",
#     "RequestId": "i2szvoUfPYBMWdSxqO3n",
#     "Status": "000"
# }
#
# possible status:
# 000 - Request initialized successfully (pending acceptance).
# 100 - The initialization request could not be completed. You can try again.
# 122 - Transaction declined to the user.
# 999 - Error on initializing the request. You can try again.

                               #  year, month, day
earlyBirdDate = datetime.datetime(2024, 10,    10)

all_prices = [
  #  non-member, member, student non-member, student member non-scholarship, student member scholarship
    [150,        90,     70,                 25,                             40                         ],
    [200,        150,    120,                60,                             70                         ]
]

def get_payment_value(user):
    """Returns the value of the payment for current user"""

    prices = all_prices[0] if earlyBirdDate > datetime.datetime.now() else all_prices[1]
    if user.student or user.unemployed:
        if user.member:
            if user.scholarship:
                return prices[4]
            else:
                return prices[3]
        else:
            return prices[2]
    else:
        if user.member:
          return prices[1]
        else:
          return prices[0]

HEADERS_JSON = {
    "accept": "application/json",
    "content-type": "application/json"
}


@bp.route('/IX_IPC/payment_mbway/<language:language>', methods=['POST'])
@login_IXIPC_required
def payment_mbway(language='pt'):
    """Starts an MBWay payment"""

    # parse and validate the phone number
    try:
        number = phonenumbers.parse(str(request.form['number']), 'PT')
        if phonenumbers.is_valid_number(number):
            # TODO: implement masking
            national_number = str(number.national_number)
            # national_number[2:-2] = re.sub('\d', '*', national_number[2:-2])
            masked = '+' + str(number.country_code) + ' ' + national_number
        else:
            raise phonenumbers.NumberParseException
    except:
        return json.dumps({'error': 'Invalid phone number'}), 400
    
    with get_session() as db_session:
        value = get_payment_value(g.IXIPC_user)

        payment = Payment(g.IXIPC_user.id, PaymentMethod.MBWay, masked, value)
        db_session.add(payment)
        db_session.commit()

        # Request to ifthenpay
        url = "https://ifthenpay.com/api/spg/payment/mbway"
            
        request_contents = {
            "mbWayKey": app.config['MBWAY_KEY'], #(string) mandatory (assigned by ifthenpay)
            "orderId": payment.transaction_id, #//(string) mandatory (15 chars max)
            "amount": str(payment.value), #(string) mandatory (decimal separator ".")
            "mobileNumber": str(number.country_code) + '#' + str(number.national_number), #(string) mandatory
            "email": "", #(string) optional (100 chars max)
            "description": app.i18n.l10n[language].format_value('IXIPC-payment-description'), #(string) optional (100 chars max)
        }

        return_content = None
        tries = 0
        while tries < 3:
            # TODO: activate again after the tests
            # response = requests.post(url, json=request_contents, headers=HEADERS_JSON)
            # response_data = response.json()

            # Fake response data for the test phase
            response_data = {
                "Amount": str(payment.value),
                "Message": "Pending",
                "OrderId": payment.transaction_id,
                "RequestId": "DxQI6XtpEPJS4BNfZVsY",
                "Status": "000"
            }
            payment.status_code = response_data['Status']
            payment.request_id = response_data['RequestId']

            if payment.status_code == '000':
                return_content = json.dumps({'id': payment.id}), 200

                user = db_session.get(User, payment.user_id)
                user.payment_id = payment.id
                break
            elif payment.status_code in {'100', '999'}:
                tries += 1 
            elif payment.status_code == '122':
                payment.failed(db_session)
                return_content = json.dumps({'error': 'Phone number declined by the payment processor'}), 400
                break
            else:
                payment.failed(db_session)
                return_content = json.dumps({'error': 'Payment failed for unknown reason'}), 400
                break
                
        if tries >= 3:
            return_content = json.dumps({'error': 'Maximum payment tries exceeded. Try again later'}), 400
            payment.failed(db_session)

        db_session.commit()

        return return_content

# https://ifthenpay.com/api/spg/payment/mbway/status?mbWayKey={mbWayKey}&requestId={requestId}
# {
#     "CreatedAt": "03-01-2024 15:15:06",
#     "Message": "Success",
#     "RequestId": "eR6mcnJzjFx7kOL1Ybdp",
#     "Status": "000",
#     "UpdateAt": "03-01-2024 15:15:16"
# }
# 000 - Transaction successfully completed (Payment confirmed).
# 020 - Transaction rejected by the user.
# 101 - Transaction expired (the user has 4 minutes to accept the payment in the MB WAY App before expiring)
# 122 - Transaction declined to the user. 

@bp.route('/IX_IPC/check_mbway_status/<language:language>', methods=['POST'])
@login_IXIPC_required
def check_mbway_status(language='pt'):
    """Checks the status of an MBWay payment"""

    with get_session() as db_session:
        payment = db_session.get(Payment, request.form['payment_id'])

        if g.IXIPC_user.id == payment.user_id:
            # return json.dumps({'status': payment.status}), 200

            # TODO: solve after dealing with the issue of not knowing what happens if the payment is waiting
            if payment.status != PaymentStatus.pending:
                return json.dumps({'status': payment.status}), 200

            url = f"https://ifthenpay.com/api/spg/payment/mbway/status?mbWayKey={app.config['MBWAY_KEY']}&requestId={payment.request_id}"

            response = requests.get(url)
            response_data = response.json()
            payment.status_code = response_data['Status']

            # TODO: answer: What if the payment is waiting??
            if payment.status_code == '000':
                payment.success(db_session)
            elif payment.status_code == '020':
                payment.canceled(db_session)
            elif payment.status_code == '101':
                payment.expired(db_session)
            elif payment.status_code == '122':
                payment.failed(db_session)
            else:
                db_session.commit()
                app.send_email(
                    'Payment failed',
                    f'Payment of user {g.IXIPC_user.id} failed for unknown reason. Error code {response_data["Status"]}, RequestId {response_data["RequestId"]}, Message {response_data["Message"]}.',
                    [app.config['MAINTENANCE_EMAIL']]
                )
                return json.dumps({'error': 'Request failed for unknown reason'}), 400
            
            db_session.commit()
            return json.dumps({
                'status': payment.status
            }), 200
        else:
            return json.dumps({'error': 'access unauthorized'}), 401


# http://www.yoursite.com/callback.php?key=[ANTI_PHISHING_KEY]&orderId=[ORDER_ID]&amount=[AMOUNT]&requestId=[REQUEST_ID]&payment_datetime=[PAYMENT_DATETIME]
@bp.route('/IX_IPC/payment_callback', methods=['GET'])
def payment_callback():
    """Callback that registers the result of an MBWay payment"""

    if request.args['key'] == app.config['ANTI_PHISHING_KEY']:
        try:
            with get_session() as db_session:
                payment = db_session.execute(
                    select(Payment).where(Payment.request_id == request.args['requestId'])
                ).scalar_one()

                if payment:
                    payment.success(db_session)
                    payment.value = request.args['amount']

                    db_session.commit()
        except:
            print(f'error: could not find payment {request.args["requestId"]} in the database')
            return '', 500

    return '', 200


@bp.route('/IX_IPC/creditcard_payment/<language:language>/start', methods=['POST'])
@login_IXIPC_required
def start_creditcard_payment(language):
    """Starts a new credit card payment and returns a user url for the payment"""

    if g.IXIPC_user.paid_registration:
        return json.dumps({'error': 'User has already paid'}), 403

    else:
        url = f'https://ifthenpay.com/api/creditcard/init/{app.config["CCARD_KEY"]}'
        value = get_payment_value(g.IXIPC_user)

        with get_session() as db_session:
            payment = Payment(g.IXIPC_user.id, PaymentMethod.Card, '', value)

            db_session.add(payment)
            db_session.commit()

            user = db_session.get(User, g.IXIPC_user.id)
            user.payment_id = payment.id

            request_data = {
                "orderId": payment.transaction_id,
                "amount": str(value),
                "successUrl": url_for('IX_IPC.creditcard_payment_success', language=language, _external=True),
                "errorUrl": url_for('IX_IPC.creditcard_payment_error', language=language, _external=True),
                "cancelUrl": url_for('IX_IPC.creditcard_payment_canceled', language=language, _external=True),
                "language" : language
            }

            return_content = None
            tries = 0
            while tries < 3:
                response = requests.post(url, json=request_data, headers=HEADERS_JSON)
                response_data = response.json()

                payment.status_code = response_data['Status']
                payment.request_id = response_data['RequestId']

                if payment.status_code == '0':
                    return_content = json.dumps({
                        'id': payment.id, 
                        'url': response_data['PaymentUrl']
                    }), 200
                    break
                elif payment.status_code == '-1':
                    tries += 1
                    print(f'error obtaining payment url for user id {g.IXIPC_user.id}: {response_data["Message"]}')
                else:
                    payment.failed(db_session)
                    return_content = json.dumps({'error': 'Unable to obtain payment link for unknown reason'}), 400
                    break
                    
            if tries >= 3:
                return_content = json.dumps({'error': 'Maximum payment tries exceeded. Try again later'}), 400
                payment.failed(db_session)

            db_session.commit()

            return return_content


def validate_payment(args):
    """Validates a payment confirmation"""

    message = str(args['id']) + str(args['amount']) + str(args['requestId'])
    hash = hmac.new(
        bytes(app.config['CCARD_KEY'], 'UTF-8'),
        message.encode(),
        sha256
    ).hexdigest()
    return hash == args['sk']


# Test cards:
# Pagamento com sucesso:
# 4012 0010 3714 1112
# CVC: 212
# Validade: 12/27

# Falha no pagamento:
# 4761 7390 0101 0135
# CVC: 608
# Validade: 12/22

@bp.route('/IX_IPC/creditcard_payment/<language:language>/success', methods=['GET'])
@login_IXIPC_required
def creditcard_payment_success(language):
    """Registers the success of a payment and redirects the user to the main page"""

    try:
        if validate_payment(request.args):
            with get_session() as db_session:
                payment = db_session.execute(
                    select(Payment).where(Payment.request_id == request.args['requestId'])
                ).scalar_one()

                payment.success(db_session)
                db_session.commit()
        else:
            print(f'Unable to validate payment {request.args}')
            flash('Could not validate credit card payment! Please contact the organizers')
    except:
        print(f'error: could not find payment {request.args["requestId"]} in the database')

    return redirect(url_for('IX_IPC.IXIPC', language=language))


@bp.route('/IX_IPC/creditcard_payment/<language:language>/canceled', methods=['GET'])
@login_IXIPC_required
def creditcard_payment_canceled(language):
    """Registers that a payment was canceled and redirects the user to the main page"""

    try:
        with get_session() as db_session:
            payment = db_session.execute(
                select(Payment).where(Payment.request_id == request.args['requestId'])
            ).scalar_one()

            print(type(payment))

            payment.canceled(db_session)
            db_session.commit()
    except:
        print(f'error: could not find payment {request.args["requestId"]} in the database')

    return redirect(url_for('IX_IPC.IXIPC', language=language))


@bp.route('/IX_IPC/creditcard_payment/<language:language>/error', methods=['GET'])
@login_IXIPC_required
def creditcard_payment_error(language):
    """Registers that a payment failed and redirects the user to the main page"""

    try:
        with get_session() as db_session:
            payment = db_session.execute(
                select(Payment).where(Payment.request_id == request.args['requestId'])
            ).scalar_one()

            payment.failed(db_session)
            db_session.commit()
    except:
        print(f'error: could not find payment {request.args["requestId"]} in the database')

    return redirect(url_for('IX_IPC.IXIPC', language=language))


@bp.route('/IX_IPC/submit_registration/<language:language>', methods=['POST'])
@login_IXIPC_required
def submit_registration(language):
    """Submit registration

    Submits a user registration to the event.
    """
    
    proof_id = request.form['proofId']

    if g.IXIPC_user.paid_registration:
        print(f'user {g.IXIPC_user.id} already paid')
        return json.dumps({'error': 'User has already paid'}), 403
    elif g.IXIPC_user.payment_id:
        print(f'user {g.IXIPC_user.id} waiting for payment validation')
        return json.dumps({'error': 'User payment being processed'}), 403
    else:
        # TODO update this function!!!!
        value = get_payment_value(g.IXIPC_user)

        with get_session() as db_session:
            payment = Payment(g.IXIPC_user.id, PaymentMethod.Transfer, proof_id, value)

            db_session.add(payment)
            db_session.commit()

            user = db_session.get(User, g.IXIPC_user.id)
            user.payment_id = payment.id

            db_session.commit()

            print(f'user {g.IXIPC_user.id} registered')

    return '', 200


@bp.route('/IX_IPC/management/logout')
@bp.route('/IX_IPC/management/logout/<language:language>')
@login_IXIPC_management_required
def management_logout(language='pt'):
    """Logs a manager out
    
    Logs a manager out and redirects to the main page
    """

    session.clear()
    return redirect(url_for('IX_IPC.IXIPC', language=language))


@bp.route('/IX_IPC/management/email/send', methods=['POST'])
@bp.route('/IX_IPC/management/email/send/<language:language>', methods=['POST'])
@login_IXIPC_management_required
def send_email(language='pt'):
    """Sends an email"""

    user_code = request.form['user']
    subject = request.form['subject']
    message = request.form['message']
    
    attachment = None
    if 'attachment' in request.form:
        attachment = request.form['attachment']

    users = []
    with get_session() as db_session:
        if user_code == 'ALL':
            users = db_session.execute(select(User)).scalars().all()
        
        elif user_code == 'ACCEPTED_ABSTRACTS':
            abstracts = db_session.execute(
                abstract_filters['accepted']
            ).scalars().all()
            for abstract in abstracts:
                user_profile = db_session.get(User, abstract.owner)
                user_profile.abstract_title = abstract.title
                user_profile.abstract_type = AbstractType.to_string(abstract.abstract_type).lower()

                users.append(user_profile)

        elif user_code == 'ALL_PAID':
            users = db_session.execute(
                select(User).filter_by(paid_registration = True)
            ).scalars().all()

        elif user_code == 'ALL_NOT_PAID':
            users = db_session.execute(
                select(User).filter_by(paid_registration = False)
            ).scalars().all()
        
        else:
            users = [db_session.get(User, int(user_code))]

    for user in users:
        first_name, last_name, name = '', '', ''
        if user.first_name:
            first_name = user.first_name if len(user.first_name) else user.name.strip().split()[0]
        else:
            first_name = user.name.strip().split()[0]

        if user.last_name:
            last_name = user.last_name
        
        if user.name:
            name = user.name

        abstract_title = ''
        abstract_type = ''
        if user_code == 'ACCEPTED_ABSTRACTS':
            abstract_title = user.abstract_title
            abstract_type = user.abstract_type

        subject_user = subject.format(first_name=first_name, last_name=last_name, name=name, email=user.email, abstract_title=abstract_title, abstract_type=abstract_type)
        message_user = message.format(first_name=first_name, last_name=last_name, name=name, email=user.email, abstract_title=abstract_title, abstract_type=abstract_type)
        app.send_email(subject_user, message_user, [user.email], attachment)

    return json.dumps(''), 200


@bp.route('/IX_IPC/management')
@bp.route('/IX_IPC/management/<language:language>')
@login_IXIPC_management_required
def management(language='pt'):
    """
    """

    return render_template(
        'management/management.html',
        lang=language,
        text_column=True
    )


@bp.route('/IX_IPC/management/participants_csv_summary')
@login_IXIPC_management_required
def participants_csv_summary():    
    with StringIO() as buffer:
        writer = csv.writer(buffer, delimiter=';')

        fields = ['first_name', 'last_name', 'email', 'institution', 'student', 'scholarship', 'unemployed', 'paid', 'submitted_talks', 'submitted_posters', 'talk_competition', 'photography_competition']
        writer.writerow(fields)

        with get_session() as db_session:
            participants = db_session.execute(
                select(User)
                  .order_by(User.first_name, User.last_name)
            ).scalars()

            for participant in participants:
                abstracts_talk = db_session.execute(
                    select(Abstract)
                      .where(Abstract.owner == participant.id)
                      .where(Abstract.abstract_type == AbstractType.PRESENTATION)
                ).scalars().all()

                abstracts_poster = db_session.execute(
                    select(Abstract)
                      .where(Abstract.owner == participant.id)
                      .where(Abstract.abstract_type == AbstractType.POSTER)
                ).scalars().all()

                writer.writerow([
                    participant.first_name, 
                    participant.last_name,
                    participant.email,
                    participant.institution,
                    1 if participant.student else 0,
                    1 if participant.scholarship else 0,
                    1 if participant.unemployed else 0,
                    1 if participant.paid_registration else 0,
                    len(abstracts_poster),
                    len(abstracts_talk),
                    1 if participant.competition_talk else 0,
                    1 if participant.competition_photography else 0
                ])

            buffer.seek(0)
            return send_file(
                BytesIO(buffer.getvalue().encode('utf-8-sig')),
                as_attachment=True,
                download_name='participants_summary.csv',
                mimetype='text/csv'
            )


@bp.route('/IX_IPC/management/participants_pdf_report')
@login_IXIPC_management_required
def participants_pdf_report():
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=A4, showBoundary=0, title='Participants list')

    w, h = A4
    mx = 3*cm
    mt = doc._topMargin
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle('Name', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=16)
    bold_style = ParagraphStyle('Bold', parent=styles['Normal'], fontName='Helvetica-Bold')

    def first_page(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 24)
        canvas.drawString(mx, mt - 5*cm, "IXIPC participants list")
        canvas.setFont("Helvetica", 14)
        canvas.drawString(mx, mt - 6*cm, f'(generated {datetime.datetime.utcnow().strftime("%d/%m/%Y, %H:%M")})')    

        canvas.setFont('Helvetica', 10)
        canvas.drawRightString(doc._rightMargin, 1.8*cm, f"{doc.page}")

        canvas.restoreState()

    def later_pages(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 10)
        canvas.drawRightString(doc._rightMargin, 1.8*cm, f"{doc.page}")
        canvas.restoreState()
    
    with get_session() as db_session:
        participants = db_session.execute(
            select(User)
              .where(User.paid_registration == True)
              .order_by(User.first_name, User.last_name)
        ).scalars()

        story = [Spacer(1, 8*cm)]#[PageBreak()]
        style = styles["Normal"]

        for participant in participants:
            participant_description = [
                Paragraph(f'{participant.first_name} {participant.last_name} {"(student)" if participant.student == True else ""}', name_style),
                Spacer(1, 0.5*cm),
                Paragraph(f'email: {participant.email}', style),
                Spacer(1, 0.2*cm),
                Paragraph(f'Institution: {participant.institution}', style),
                Spacer(1, 0.2*cm),
            ]

            if participant.paid_registration and participant.payment_id:
                payment = db_session.get(Payment, participant.payment_id)
                method = PaymentMethod.to_str(payment.method)
                participant_description.append(Paragraph(
                    f'Paid: {payment.value}€ on {payment.concluded.strftime("%d/%m/%Y")} (via {method})',
                    style)
                )
            else:
                participant_description.append(Paragraph('No payment registered', style))
            participant_description.append(Spacer(1, 0.4*cm))

            submitted_abstracts = db_session.execute(
                select(Abstract)
                  .where(Abstract.owner == participant.id)
                  .where(Abstract.submitted == True)
            ).scalars().all()

            if submitted_abstracts:
                participant_description.append(
                    Paragraph('Submitted abstracts:', bold_style)
                )
                participant_description.append(Spacer(1, 0.2*cm))

                for i, abstract in enumerate(submitted_abstracts):
                    participant_description.append(
                        Paragraph(
                            f'{i+1}. ' 
                            + AbstractType.to_string(abstract.abstract_type) 
                            + ' — ' + abstract.title, 
                        style)
                    )

                    authors = db_session.execute(
                        select(AbstractAuthor)
                          .where(AbstractAuthor.abstract_id == abstract.id)
                          .order_by(AbstractAuthor.order)
                    ).scalars().all()

                    authors_list = ''
                    for author in authors:
                        author_details = db_session.get(Author, author.author_id)
                        if author.presenter:
                            authors_list += f'<u>{author_details.first_name} {author_details.last_name}</u>, '
                        else:
                            authors_list += f'{author_details.first_name} {author_details.last_name}, '

                    participant_description.append(Spacer(1, 0.2*cm))
                    participant_description.append(
                        Paragraph(f'Authors: ' + authors_list[:-2], style)
                    )

                    participant_description.append(Spacer(1, 0.2*cm))
                    participant_description.append(
                        Paragraph(f'Keywords: ' + abstract.keywords, style)
                    )
                    participant_description.append(Spacer(1, 0.4*cm))
            else:
                participant_description.append(
                    Paragraph('The participant did not submit any abstracts', 
                    style)
                )
            participant_description.append(Spacer(1, 0.2*cm))

            story.append(KeepTogether(participant_description))
            story.append(Spacer(1, 1*cm))

        doc.build(story, onFirstPage=first_page, onLaterPages=later_pages)

    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name='participants_list.pdf',
        mimetype='application/pdf'
    )


@bp.route('/IX_IPC/management/participants')
@bp.route('/IX_IPC/management/participants/<language:language>')
@login_IXIPC_management_required
def participants_list(language='pt'):
    """Gets a list of participants"""

    # TODO: include option to limit to paid registrations
    with get_session() as db_session:
        participants = db_session.execute(
            select(User).order_by(User.first_name, User.last_name)
        ).scalars()

        parts = []
        for participant in participants:
            if participant.paid_registration and participant.payment_id:
                participant.payment = db_session.get(Payment, participant.payment_id)

            participant.submitted_abstracts = db_session.execute(
                select(Abstract)
                  .where(Abstract.owner == participant.id)
                  .where(Abstract.submitted == True)
            ).scalars().all()

            parts.append(participant)

        return render_template(
            'management/participant_list.html',
            lang=language,
            text_column=True,
            participants=parts
        )


@bp.route('/IX_IPC/management/dinner')
@bp.route('/IX_IPC/management/dinner/<language:language>')
@login_IXIPC_management_required
def dinner_list(language='pt'):
    """Shows the dinner information"""

    with get_session() as db_session:
        return render_template(
            'management/dinner_details.html',
            lang=language,
            text_column=True,
        )

@bp.route('/IX_IPC/management/abstracts')
@bp.route('/IX_IPC/management/abstracts/<language:language>')
@login_IXIPC_management_required
def abstracts_list(language='pt'):
    """Shows the list of abstracts"""

    with get_session() as db_session:
        return render_template(
            'management/abstracts_list.html',
            lang=language,
            text_column=True,
        )


@bp.route('/IX_IPC/management/load_abstracts_list', methods=['POST'])
@login_IXIPC_management_required
def load_abstracts_list():
    """Loads the list of abstracts"""

    with get_session() as db_session:
        abstracts_list = db_session.execute(
            select(Abstract)
        ).scalars().all()

        abstracts = []
        for abstract in abstracts_list:
            created_by = db_session.get(User, abstract.owner)
            abstract = {
                'id': abstract.id,
                'href': url_for('IX_IPC.abstract_details', id=abstract.id),
                'type': AbstractType.to_string(abstract.abstract_type),
                'title': abstract.title,
                'abstract': abstract.abstract,
                'keywords': abstract.keywords,
                'submitted': abstract.submitted, 
                'acceptance_status': abstract.acceptance_status,
                'created_by': f'{created_by.first_name} {created_by.last_name} ({created_by.email})'
            }
            abstracts.append(abstract)

        return json.dumps({
            'abstracts': abstracts
        }), 200


@bp.route('/IX_IPC/management/participant/<int:id>')
@bp.route('/IX_IPC/management/participant/<int:id>/<language:language>')
@login_IXIPC_management_required
def participant_details(id, language='pt'):
    with get_session() as db_session:
        participant = db_session.get(User, id)

        participant.submitted_abstracts = db_session.execute(
            select(Abstract)
              .where(Abstract.owner == participant.id)
              .where(Abstract.submitted == True)
        ).scalars().all()

        if participant.paid_registration and participant.payment_id:
            participant.payment = db_session.get(Payment, participant.payment_id)

        return render_template(
            'management/participant_details.html',
            lang=language,
            text_column=True,
            participant=participant
        )


@bp.route('/IX_IPC/management/fetch_participant_emails', methods=['POST'])
@login_IXIPC_management_required
def fetch_participant_emails():
    """Fetches the emails list of participant"""

    with get_session() as db_session:
        user = db_session.get(User, request.form['user_id'])

        emails = []            
        email_list = app.fetch_user_emails(user.email)
        for type_ in ['received', 'sent']:
            for email in email_list[type_]:
                emails.append({
                    'date': email.date.strftime("%d/%m/%Y (%H:%M:%S)"),
                    'timestamp': str(email.date),
                    'subject': email.subject,
                    'message': email.text,
                    'type': type_.capitalize(),
                    'attachments': [f'{file.filename} ({int(file.size * 0.0009765625)} KB)' for file in email.attachments]
                })
        emails.sort(key=lambda email: email['timestamp'], reverse=True)

        return json.dumps({
            'emails': emails
        }), 200


@bp.route('/IX_IPC/management/update_abstract_acceptance_status/<int:id>/<int:new_status>')
@login_IXIPC_management_required
def update_abstract_acceptance_status(id, new_status):
    with get_session() as db_session:
        try:
            abstract = db_session.get(Abstract, id)
            if new_status == 0:
                abstract.undecide()
            elif new_status == 1:
                abstract.accept()
            elif new_status == 3:
                abstract.conditionally_accept()
            else:
                abstract.reject()

            db_session.commit()
            return json.dumps(''), 200
        except:
            return json.dumps({
                'error': f'Failed to update abstract {id}.'
            }), 500


@bp.route('/IX_IPC/management/unsubmit_abstract/<int:id>')
@login_IXIPC_management_required
def force_unsubmit(id):
    """Unsubmit a previously submitted abstract"""
    with get_session() as db_session:
        try:
            abstract = db_session.get(Abstract, id)
            abstract.unsubmit()
            db_session.commit()
            return json.dumps(''), 200
        except:
            return json.dumps({
                'error': f'Failed to unsubmit abstract {id}.'
            }), 500

abstract_filters = {
    'all': select(Abstract).order_by(Abstract.id),
    'submitted': select(Abstract)
                  .where(Abstract.submitted == True)
                  .order_by(Abstract.id),
    'undecided': select(Abstract)
                  .where(Abstract.submitted == True)
                  .where(Abstract.acceptance_status == 0)
                  .order_by(Abstract.id),
    'accepted':  select(Abstract)
                  .where(Abstract.submitted == True)
                  .where(Abstract.acceptance_status == 1)
                  .order_by(Abstract.id),
    'rejected':  select(Abstract)
                  .where(Abstract.submitted == True)
                  .where(Abstract.acceptance_status == 2)
                  .order_by(Abstract.id),
}

def process_authors_affiliations(abstract_id):
    author_list = []
    affiliations_list = []

    with get_session() as db_session:
        authors = db_session.execute(
            select(AbstractAuthor)
              .where(AbstractAuthor.abstract_id == abstract_id)
              .order_by(AbstractAuthor.order)
        ).scalars()

        for author in authors:
            author.details = db_session.get(Author, author.author_id)
            author.affiliations = Affiliation.get_list(db_session, author.author_id)

            author_desc = {
                'first_name': author.details.first_name,
                'last_name': author.details.last_name,
                'presenter': author.presenter,
                'formatted': '',
                'affiliations': []
            }

            if author.presenter:
                author_desc['formatted'] += f'<u>{author.details.first_name} {author.details.last_name}</u>'
            else:
                author_desc['formatted'] += f'{author.details.first_name} {author.details.last_name}'


            for affiliation in author.affiliations:
                id = affiliation.institution_id

                for i, registered_affiliation in enumerate(affiliations_list):
                    if id == registered_affiliation.id:
                        author_desc['affiliations'].append(i)
                        break
                else:
                    author_desc['affiliations'].append(len(affiliations_list))
                    affiliations_list.append(affiliation.institution)
                    
            author_list.append(author_desc)

    return author_list, affiliations_list


import shutil
import os
from glob import glob

@bp.route('/IX_IPC/management/docx_list/<string:filter>')
@bp.route('/IX_IPC/management/docx_list')
@login_IXIPC_management_required
def docx_list(filter=''):
    try:
        os.remove(f"{app.config['TEMP_FOLDER']}/abstracts.zip")
    except:
        pass

    for f in glob(f"{app.config['TEMP_FOLDER']}/abstracts/*.docx"):
        try:
            os.remove(f)
        except:
            pass
    
    file_list = []
    with get_session() as db_session:
        if filter not in abstract_filters.keys():
            return json.dumps({'error', 'unrecognized filter'}), 500
        
        abstracts = db_session.execute(abstract_filters[filter]).scalars()

        for i, abstract in enumerate(abstracts):
            file_list.append(generate_docx_abstract(abstract, i+1))
    
    shutil.make_archive(
        f"{app.config['TEMP_FOLDER']}/abstracts",
        'zip',
        f"{app.config['TEMP_FOLDER']}/abstracts/"
    )

    return send_from_directory(app.config['TEMP_FOLDER'], 'abstracts.zip')


def generate_docx_abstract(abstract, j):
    with get_session() as db_session:
        authors, affiliations = process_authors_affiliations(abstract.id)
        owner = db_session.get(User, abstract.owner)
        
        owner_name = ''
        if owner.first_name and owner.last_name:
            owner_name = f"{owner.first_name.lower().replace(' ', '_')}_{owner.last_name.lower().replace(' ', '_')}"
        else:
            owner_name = 'no_name'
        
        abstract_type = AbstractType.to_string(abstract.abstract_type).lower().replace(' ', '_')

        folder = f"{app.config['TEMP_FOLDER']}/abstracts/"
        filename = f"{abstract_type}_{owner_name}_{j}.docx"
        
        doc = Document()
        paragraph_style = doc.styles['Normal']
        paragraph_style.font.name = 'Helvetica' 

        doc.add_heading(f'{AbstractType.to_string(abstract.abstract_type)} — {abstract.title}', 0)
        
        p = doc.add_paragraph()
        p.add_run('Scientific area: ').bold = True
        p.add_run(abstract.scientific_area.lower())

        p = doc.add_paragraph()
        p.add_run('Abstract: ').bold = True
        p.add_run(abstract.abstract)

        p = doc.add_paragraph()
        p.add_run('Authors: ').bold = True
        p.paragraph_format.space_after = Pt(1)

        p = doc.add_paragraph()
        for k, author in enumerate(authors):
            author_run = p.add_run(f"{author['first_name']} {author['last_name']} ")
            if author['presenter']:
                author_run.underline = True

            author_affiliations = ''
            for i in author['affiliations']:
                author_affiliations += f'{i+1},'
            p.add_run(author_affiliations[:-1]).font.superscript = True

            if k + 1 < len(authors):
                p.add_run(', ')
        p.paragraph_format.space_after = Pt(3)

        for i, affiliation in enumerate(affiliations):
            p = doc.add_paragraph()
            r = p.add_run(f"[{i+1}] {affiliation.name}, {affiliation.address} ({affiliation.country.capitalize()})")
            r.italic = True
            r.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_after = Pt(11)  

        p = doc.add_paragraph()
        p.add_run('Keywords: ').bold = True
        p.add_run(abstract.keywords)

        p = doc.add_paragraph()
        p.add_run(f"Created by {owner.first_name} {owner.last_name} ({owner.email})")
        p.paragraph_format.space_after = Pt(36)

        doc.add_paragraph(f"Submitted on {abstract.submitted_on.strftime('%d/%m/%Y')}").paragraph_format.space_after = Pt(3)
        doc.add_paragraph(f"Docx generated on {datetime.datetime.now().strftime('%d/%m/%Y')}")

        doc.save(folder + filename)
        return filename


@bp.route('/IX_IPC/management/abstracts_pdf_report')
@bp.route('/IX_IPC/management/abstracts_pdf_report/<string:filter>')
@login_IXIPC_management_required
def abstracts_pdf_report(filter=''):
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=A4, showBoundary=0, title='Abstracts list')

    w, h = A4
    mx = 3*cm
    mt = doc._topMargin
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle('Name', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=16, leading=16)
    bold_style = ParagraphStyle('Bold', parent=styles['Normal'], fontName='Helvetica-Bold')
    style_small = ParagraphStyle('Small', parent=styles['Normal'], fontSize=8)

    def first_page(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 24)
        canvas.drawString(mx, mt - 5*cm, f"IXIPC abstracts list ({filter})")
        canvas.setFont("Helvetica", 14)
        canvas.drawString(mx, mt - 6*cm, f'(generated {datetime.datetime.utcnow().strftime("%d/%m/%Y, %H:%M")})')    

        canvas.setFont('Helvetica', 10)
        canvas.drawRightString(doc._rightMargin, 1.8*cm, f"{doc.page}")

        canvas.restoreState()

    def later_pages(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 10)
        canvas.drawRightString(doc._rightMargin, 1.8*cm, f"{doc.page}")
        canvas.restoreState()
    
    with get_session() as db_session:
        if filter not in abstract_filters.keys():
            return json.dumps({'error', 'unrecognized filter'}), 500
        
        abstracts = db_session.execute(abstract_filters[filter]).scalars()

        story = [Spacer(1, 8*cm)] #[PageBreak()]
        style = ParagraphStyle('BaseStyle', parent=styles['Normal'], alignment=4)

        for i, abstract in enumerate(abstracts):
            abstract_description = [
                Paragraph(f'{i+1}. {AbstractType.to_string(abstract.abstract_type)} — {abstract.title}', name_style),
                Spacer(1, 0.5*cm),
                Paragraph(f'<b>Scientific area:</b> {abstract.scientific_area.lower()}', style),
                Spacer(1, 0.2*cm),
                Paragraph(f'<b>Abstract:</b> {abstract.abstract}', style),
                Spacer(1, 0.2*cm),
                Paragraph(f'<b>Authors:</b>', style),
                Spacer(1, 0.05*cm)
            ]

            authors, affiliations = process_authors_affiliations(abstract.id)

            author_str = ''
            for author in authors:
                author_str += author['formatted'] + '<sup>'
                for i in author['affiliations']:
                    author_str += f'{i+1},'
                author_str = author_str[:-1] + '</sup>, '

            abstract_description.append(Paragraph(author_str[:-2], style))
            abstract_description.append(Spacer(1, 0.1*cm))

            for i, affiliation in enumerate(affiliations):
                abstract_description.append(Paragraph(
                    f'  <i>[{i+1}] {affiliation.name}, {affiliation.address} ({affiliation.country.capitalize()})</i>', 
                    style_small
                ))
                abstract_description.append(Spacer(1, 0.05*cm))
                    
            abstract_description.append(Spacer(1, 0.4*cm))
            abstract_description.append(
                Paragraph(f'<b>Keywords:</b> ' + abstract.keywords, style)
            )
            abstract_description.append(Spacer(1, 0.4*cm))

            owner = db_session.get(User, abstract.owner)
            abstract_description.append(
                Paragraph(f'Created by {owner.first_name} {owner.last_name} ({owner.email})')
            )
            abstract_description.append(Spacer(1, 0.2*cm))

            try:
                abstract_description.append(
                    Paragraph(f'Submitted on {abstract.submitted_on.strftime("%d/%m/%Y")}')
                )
                abstract_description.append(Spacer(1, 0.2*cm))
            except:
                pass
            
            story.append(KeepTogether(abstract_description))
            story.append(Spacer(1, 1*cm))
            
        doc.build(story, onFirstPage=first_page, onLaterPages=later_pages)

    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name='abstracts_list.pdf',
        mimetype='application/pdf'
    )


@bp.route('/IX_IPC/management/abstracts_csv_summary')
@login_IXIPC_management_required
def abstracts_csv_summary():    
    with StringIO() as buffer:
        writer = csv.writer(buffer, delimiter=';')

        fields = ['type', 'title', 'abstract', 'keywords', 'submitted_by', 'submitted_by_email', 'presenter', 'presenter_email']
        writer.writerow(fields)

        with get_session() as db_session:
            abstracts = db_session.execute(
                select(Abstract)
                  .where(Abstract.submitted == True)
                  .order_by(Abstract.id)
            ).scalars()

            for abstract in abstracts:
                user = db_session.get(User, abstract.owner)

                presenter_info = db_session.execute(
                    select(AbstractAuthor)
                      .where(AbstractAuthor.abstract_id == abstract.id)
                      .where(AbstractAuthor.presenter == True)
                ).scalars().all()

                class Object(object):
                    pass
                
                presenter = Object()
                if len(presenter_info):
                    presenter = db_session.get(Author, presenter_info[0].author_id)
                else:
                    presenter.first_name = 'unknown'
                    presenter.last_name = ''
                    presenter.email = ''

                writer.writerow([
                    abstract.abstract_type, 
                    abstract.title,
                    abstract.abstract,
                    f'\"{abstract.keywords}\"',
                    f'{user.first_name} {user.last_name}',
                    user.email,
                    f'{presenter.first_name} {presenter.last_name}',
                    presenter.email,
                ])

            buffer.seek(0)
            return send_file(
                BytesIO(buffer.getvalue().encode('utf-8-sig')),
                as_attachment=True,
                download_name='abstracts_summary.csv',
                mimetype='text/csv'
            )


@bp.route('/IX_IPC/management/abstract_details/<int:id>')
@bp.route('/IX_IPC/management/abstract_details/<int:id>/<language:language>')
@login_IXIPC_management_required
def abstract_details(id, language='pt'):
    with get_session() as db_session:
        abstract = db_session.get(Abstract, id)

        abstract.user = db_session.get(User, abstract.owner)
        abstract.authors, abstract.affiliations = process_authors_affiliations(abstract.id)

        for author in abstract.authors:
            author['affiliations'] = map(lambda x: str(x+1), author['affiliations'])

        return render_template(
            'management/abstract_details.html',
            abstract=abstract,
            lang=language,
            text_column=True
        )


@bp.route('/IX_IPC/management/registrations/')
@bp.route('/IX_IPC/management/registrations/<language:language>')
@login_IXIPC_management_required
def registrations(language='pt'):
    with get_session() as db_session:
        participants = db_session.execute(
                select(User)
                  .where(User.payment_id != None)
            ).scalars().all()
        
        for participant in participants:
            participant.payment = db_session.get(Payment, participant.payment_id)
            participant.payment.filepath = participant.first_name + participant.last_name + 'PaymentProof'
            participant.payment.filepath = ''.join(participant.payment.filepath.split())

        return render_template(
            'management/registrations.html',
            participants=participants,
            lang=language,
            text_column=True
        )


@bp.route('/IX_IPC/management/toggle_registration_status/<string:id>')
@login_IXIPC_management_required
def update_registration_status(id):
    with get_session() as db_session:
        participant = db_session.get(User, id)
        participant.paid_registration = not participant.paid_registration
        db_session.commit()

        feedback = f'user {id} registration updated ({"verified" if participant.paid_registration else "not verified"})'
        print(feedback)

        return feedback, 200


def register(app) -> None:
    """Registers the IXIPC blueprint with the Flask app"""

    app.register_blueprint(bp)

    def abstract_type_filter(abstract_type, language) -> str:
        abstract = ''
        if abstract_type == 1:
            abstract = app.i18n.l10n[language].format_value('IXIPC-abstract-poster')
        else:
            abstract = app.i18n.l10n[language].format_value('IXIPC-abstract-presentation')
        
        return abstract

    app.jinja_env.filters['abstract_type'] = abstract_type_filter

    from .db_IXIPC import init_IXIPC_db_command, set_IXIPC_organizer
    app.cli.add_command(init_IXIPC_db_command)
    app.cli.add_command(set_IXIPC_organizer)
